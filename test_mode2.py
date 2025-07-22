import docx
import msoffcrypto
from cryptography.hazmat.primitives import serialization, hashes
from cryptography.hazmat.primitives.asymmetric import padding, rsa
import base64
import subprocess
import json
import os

# --- ตั้งค่าไฟล์และพารามิเตอร์ ---
docx_filename = "secret.docx"
protected_docx = "protected.docx"
decrypted_docx = "unlocked_secret.docx"

video_input = r"C:\Users\65011\Desktop\Segano\work00002\vdio\mm.mp4"
video_output = r"C:\Users\65011\Desktop\Segano\work00002\vdio\output_with_metadata.mp4"
metadata_key = "comment"

private_key_file = "private_key.pem"  # บันทึก private key เพื่อใช้ในการถอดรหัส
text_to_write = "นี่คือข้อความลับสุดยอด 🚨"


def encrypt_and_embed():
    print("[🚀] เริ่มกระบวนการเข้ารหัสและซ่อนข้อมูล...\n")

    # ===== 1. สร้าง .docx และเขียนข้อความ =====
    doc = docx.Document()
    doc.add_paragraph(text_to_write)
    doc.save(docx_filename)
    print(f"[+] สร้างไฟล์ {docx_filename} เรียบร้อย")

    # ===== 2. ใส่รหัสผ่านให้ .docx ด้วย msoffcrypto =====
    docx_password = "Mikkee"
    with open(docx_filename, "rb") as f_in:
        with open(protected_docx, "wb") as f_out:
            office_file = msoffcrypto.OfficeFile(f_in)
            office_file.encrypt(docx_password, f_out)
    print(f"[+] ป้องกันไฟล์ด้วยรหัสผ่านแล้ว → {protected_docx}")

    # ===== 3. สร้าง RSA key และเข้ารหัสรหัสผ่าน =====
    private_key = rsa.generate_private_key(public_exponent=65537, key_size=2048)
    public_key = private_key.public_key()

    # บันทึก private key ลงไฟล์เพื่อใช้ภายหลัง
    with open(private_key_file, "wb") as key_file:
        key_file.write(
            private_key.private_bytes(
                encoding=serialization.Encoding.PEM,
                format=serialization.PrivateFormat.PKCS8,
                encryption_algorithm=serialization.NoEncryption()
            )
        )
    print(f"[+] บันทึก private key ลงไฟล์ → {private_key_file}")

    encrypted_password = public_key.encrypt(
        docx_password.encode(),
        padding.OAEP(
            mgf=padding.MGF1(algorithm=hashes.SHA256()),
            algorithm=hashes.SHA256(),
            label=None
        )
    )
    encoded_encrypted_password = base64.b64encode(encrypted_password).decode()
    print(f"[+] รหัสผ่านถูกเข้ารหัสด้วย RSA แล้ว (base64):\n{encoded_encrypted_password}")

    # ===== 4. ฝังข้อความใน metadata ของวิดีโอ =====
    try:
        subprocess.run([
            "ffmpeg", "-i", video_input,
            "-metadata", f"{metadata_key}={encoded_encrypted_password}",
            "-codec", "copy", video_output
        ], check=True)
        print(f"[+] ฝังข้อความเข้ารหัสลงใน metadata ของวิดีโอแล้ว → {video_output}")
    except subprocess.CalledProcessError as e:
        print(f"[-] เกิดข้อผิดพลาดขณะรัน ffmpeg: {e}")
        exit(1)
    except FileNotFoundError:
        print("[-] ไม่พบ ffmpeg กรุณาติดตั้งและเพิ่มลง PATH")
        exit(1)

    # ===== ล้างไฟล์ชั่วคราว =====
    if os.path.exists(docx_filename):
        os.remove(docx_filename)
        print(f"[🗑️] ลบไฟล์ชั่วคราว: {docx_filename}")


def decrypt_and_extract():
    print("[🔓] เริ่มกระบวนการถอดรหัสและดึงข้อมูล...\n")

    # ===== 1. ดึง metadata จากวิดีโอ =====
    try:
        result = subprocess.run([
            "ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", video_output
        ], capture_output=True, text=True, check=True)
        metadata = json.loads(result.stdout)
        encoded_encrypted_password = metadata['format']['tags'][metadata_key]
        print(f"[+] ดึงข้อมูลจาก metadata สำเร็จ: {encoded_encrypted_password}")
    except subprocess.CalledProcessError as e:
        print(f"[-] เกิดข้อผิดพลาดขณะรัน ffprobe: {e}")
        exit(1)
    except KeyError:
        print(f"[-] ไม่พบ metadata ที่ชื่อ '{metadata_key}' ในวิดีโอ")
        exit(1)
    except FileNotFoundError:
        print("[-] ไม่พบ ffprobe กรุณาติดตั้ง ffmpeg และเพิ่มลง PATH")
        exit(1)

    # ===== 2. อ่าน private key จากไฟล์ =====
    if not os.path.exists(private_key_file):
        print(f"[-] ไม่พบไฟล์ private key: {private_key_file}")
        print("[-] กรุณาเรียกใช้โหมดเข้ารหัสก่อน หรือวางไฟล์ private key ให้ถูกต้อง")
        exit(1)

    with open(private_key_file, "rb") as key_file:
        private_key = serialization.load_pem_private_key(
            key_file.read(),
            password=None
        )

    # ===== 3. ถอดรหัสรหัสผ่าน =====
    try:
        encrypted_password = base64.b64decode(encoded_encrypted_password)
        decrypted_password = private_key.decrypt(
            encrypted_password,
            padding.OAEP(
                mgf=padding.MGF1(algorithm=hashes.SHA256()),
                algorithm=hashes.SHA256(),
                label=None
            )
        )
        docx_password = decrypted_password.decode()
        print(f"[+] ถอดรหัสรหัสผ่านสำเร็จ: {docx_password}")
    except Exception as e:
        print(f"[-] ถอดรหัสไม่สำเร็จ: {e}")
        exit(1)

    # ===== 4. ถอดรหัส protected.docx =====
    if not os.path.exists(protected_docx):
        print(f"[-] ไม่พบไฟล์: {protected_docx}")
        exit(1)

    try:
        with open(protected_docx, "rb") as f_in:
            with open(decrypted_docx, "wb") as f_out:
                office_file = msoffcrypto.OfficeFile(f_in)
                office_file.load_key(password=docx_password)
                office_file.decrypt(f_out)  # <-- เขียนไฟล์ถอดรหัสลง f_out โดยตรง
        print(f"[+] ถอดรหัสไฟล์ protected.docx แล้วบันทึกลง → {decrypted_docx}")

        # อ่านเนื้อหา
        doc = docx.Document(decrypted_docx)
        content = "\n".join([p.text for p in doc.paragraphs])
        print(f"[📜] เนื้อหาภายในไฟล์: {content}")

    except Exception as e:
        print(f"[-] ไม่สามารถเปิดไฟล์ protected.docx ได้: {e}")
        exit(1)



# ============ เริ่มต้นโปรแกรม ============
if __name__ == "__main__":
    # ตรวจสอบว่ามี private key หรือยัง
    if os.path.exists(private_key_file):
        print("[ℹ️] พบรหัส private key → เข้าสู่โหมดถอดรหัส")
        decrypt_and_extract()
    else:
        print("[ℹ️] ไม่พบรหัส private key → เข้าสู่โหมดเข้ารหัส")
        encrypt_and_embed()